import matplotlib.pyplot as plt
from docx.shared import Inches
import os
import uuid


def embed_graph_in_doc(target, fig, width_inches=5):
    """Generic function to embed any matplotlib figure in a Word doc or cell."""
    filename = f"temp_graph_{uuid.uuid4().hex[:8]}.png"

    fig.savefig(filename, dpi=300, bbox_inches='tight')
    plt.close(fig)

    if hasattr(target, "add_picture"):
        target.add_picture(filename, width=Inches(width_inches))
    else:
        paragraph = target.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(filename, width=Inches(width_inches))

    os.remove(filename)
