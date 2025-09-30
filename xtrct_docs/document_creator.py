from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


def _remove_table_borders(table):
    """Strip borders from a table so wrapper blocks stay invisible."""
    tbl = table._element
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = tbl_borders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tbl_borders.append(border)
        border.set(qn('w:val'), 'nil')


def _prevent_row_split(row):
    """Ensure the entire row moves to a new page instead of splitting."""
    tr = row._element
    tr_pr = tr.find(qn('w:trPr'))
    if tr_pr is None:
        tr_pr = OxmlElement('w:trPr')
        tr.insert(0, tr_pr)
    cant_split = tr_pr.find(qn('w:cantSplit'))
    if cant_split is None:
        cant_split = OxmlElement('w:cantSplit')
        tr_pr.append(cant_split)


def _set_table_cant_split(table):
    for row in table.rows:
        _prevent_row_split(row)


def _chain_table_rows(table):
    """Link every row so paragraph-level keep settings cascade."""
    last_idx = len(table.rows) - 1
    for idx, row in enumerate(table.rows):
        keep_with_next = idx != last_idx
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = keep_with_next


def _prepare_table_block(table):
    _set_table_cant_split(table)
    _chain_table_rows(table)


def _clear_cell(cell):
    while cell.paragraphs:
        p = cell.paragraphs[0]._element
        cell._tc.remove(p)


def _start_question_block(doc):
    """Create a 1x1 wrapper table so an entire question stays on a page."""
    container = doc.add_table(rows=1, cols=1)
    container.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(container)
    _prevent_row_split(container.rows[0])
    cell = container.cell(0, 0)
    _clear_cell(cell)
    cell.add_paragraph()
    return container, cell

def _get_or_create_tc_pr(cell):
    tc = cell._tc
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        tc.insert(0, tc_pr)
    return tc_pr


def _shade_cell(cell, fill="E6E6E6"):
    tc_pr = _get_or_create_tc_pr(cell)
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)


def _style_button_cell(cell):
    _shade_cell(cell)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.bold = True


def _add_button_group(target_cell, unit_label, options, total_width_inches=None):
    if not options:
        return

    label_para = target_cell.add_paragraph(unit_label)
    label_para.paragraph_format.space_before = Pt(0)
    label_para.paragraph_format.keep_with_next = True
    label_para.paragraph_format.space_after = Pt(2)
    for run in label_para.runs:
        run.bold = True

    button_table = target_cell.add_table(rows=1, cols=len(options))
    button_table.style = 'Table Grid'
    _prepare_table_block(button_table)

    available_width = total_width_inches if total_width_inches is not None else 2.4
    min_per_option = 0.7
    max_per_option = 1.1
    per_option = available_width / max(1, len(options))
    per_option = max(min_per_option, min(max_per_option, per_option))
    available_width = per_option * max(1, len(options))
    column_width = Inches(per_option)
    for idx, option in enumerate(options):
        button_cell = button_table.cell(0, idx)
        button_cell.width = column_width
        button_cell.text = option
        _style_button_cell(button_cell)



def _create_side_by_side_graph_layout(container_cell, graph_data, button_options, units, graph_width_inches=None):
    if graph_data is None:
        return

    layout_table = container_cell.add_table(rows=1, cols=2)
    layout_table.autofit = False
    _remove_table_borders(layout_table)
    _prepare_table_block(layout_table)

    left_cell = layout_table.cell(0, 0)
    right_cell = layout_table.cell(0, 1)
    _clear_cell(left_cell)
    _clear_cell(right_cell)

    graph_width = graph_width_inches or 2.9
    total_width = 6.2
    side_padding = 0.3
    min_right = 2.4

    options_map = button_options or {}
    max_option_count = max((len(opts or []) for opts in options_map.values()), default=1)
    min_button_cell = 0.75
    required_right = max(min_right, max_option_count * min_button_cell + 0.5)
    max_graph = max(1.6, total_width - required_right - side_padding)
    if graph_width > max_graph:
        graph_width = max_graph

    left_width = graph_width + side_padding
    right_width = total_width - left_width
    if right_width < required_right:
        right_width = required_right
        left_width = total_width - right_width
        graph_width = max(1.5, left_width - side_padding)

    left_cell.width = Inches(left_width)
    right_cell.width = Inches(right_width)

    content_width = max(right_width - 0.6, max_option_count * min_button_cell)
    from utils.graph_utils import embed_graph_in_doc
    embed_graph_in_doc(left_cell, graph_data, width_inches=graph_width)
    for idx, unit in enumerate(units or []):
        _add_button_group(right_cell, unit, options_map.get(idx), total_width_inches=content_width)

    if not (units and options_map):
        right_cell.add_paragraph()


# Create main test document
 # Store answers as you generate
def create_vert_answer_table(target, problem_units, problem_number):
    """
    Create a formatted answer table for multipart questions
    
    :param target: Document or cell that receives the content
    :param problem_units: List of unit strings (e.g., ["Direction", "Motion State"])
    :param problem_number: The problem number for labeling
    """
    spacer_para = target.add_paragraph()
    spacer_para.paragraph_format.keep_with_next = True
    spacer_para.paragraph_format.keep_together = True

    final_answers_para = target.add_paragraph()
    final_answers_para.paragraph_format.keep_with_next = True
    final_answers_para.paragraph_format.keep_together = True
    final_answers_run = final_answers_para.add_run("Final Answers:")
    final_answers_run.bold = True

    table = target.add_table(rows=len(problem_units), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, unit in enumerate(problem_units):
        clean_unit = unit.replace('(', '').replace(')', '').strip()

        left_cell = table.cell(i, 0)
        left_cell.text = clean_unit
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        right_cell = table.cell(i, 1)
        right_cell.text = "_" * 20
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(3.0)

    _prepare_table_block(table)


def create_answer_table(target, problem_units, problem_number):
    """
    Create a formatted answer table for multipart questions with horizontal layout
    
    :param target: Document or cell that receives the content
    :param problem_units: List of unit strings (e.g., ["Direction", "Motion State"])
    :param problem_number: The problem number for labeling
    """
    final_answers_para = target.add_paragraph()
    final_answers_para.paragraph_format.keep_with_next = True
    final_answers_para.paragraph_format.keep_together = True
    final_answers_run = final_answers_para.add_run("Final Answers:")
    final_answers_run.bold = True

    num_parts = len(problem_units)
    table = target.add_table(rows=2, cols=num_parts)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, unit in enumerate(problem_units):
        header_cell = table.cell(0, i)
        header_cell.text = unit
        header_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in header_cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for i in range(num_parts):
        answer_cell = table.cell(1, i)
        answer_cell.text = ""
        answer_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    column_width = Inches(6.0 / num_parts)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = column_width

    table.rows[1].height = Inches(0.5)

    _prepare_table_block(table)


def keep_question_together(doc, question_paragraphs):
    """Apply 'keep with next' to all paragraphs in a question except the last one"""
    for i in range(len(question_paragraphs) - 1):
        question_paragraphs[i].paragraph_format.keep_with_next = True


def is_multipart_question(problem):
    """
       Determine if a question has multiple answer parts
    
    :param problem: Problem dictionary with 'answers' and 'units' keys
    :return: Boolean indicating if question is multipart
    """
    return len(problem["answers"]) > 1


def create_and_embed_graph(doc, graph_data, filename = "temp_graph.png"):
    import matplotlib.pyplot as plt
    plt.figure(figsize=(4,4))
    plt.plot(graph_data['x'], graph_data['y'])
    plt.xlabel(graph_data['xlabel'])
    plt.ylabel(graph_data['ylabel'])
    plt.title(graph_data['title'])
    plt.grid(True)

    plt.savefig(filename, dpi=300,bbox_inches='tight')
    plt.close()

    doc.add_picture(filename, width = Inches(5))

    os.remove(filename)
    

def create_doc(title: str, question_generator, number_of_docs: int, tables: bool = True):
  doc = Document()
  answer_key = [] 
  for doc_num in range(1, number_of_docs + 1):
      if doc_num != 1:
          doc.add_section(WD_SECTION_START.ODD_PAGE)
      questions = question_generator()
      header = doc.add_heading(f'{title} {doc_num}', 0)
      header.paragraph_format.keep_with_next = True
      name_line = doc.add_paragraph(f'Name: _____________________________________________ Date: __________________')
      name_line.paragraph_format.keep_with_next = True
      
      problem_number = 1
      version_answers = {}
      section_num = 1

      for section in questions:
         heading = section["heading"]
         section_heading = doc.add_heading(f"Section {section_num}: {heading}", level=2)
         section_heading.paragraph_format.keep_with_next = True
         instructions = section.get("section_instructions")
         if instructions:
             instructions_para = doc.add_paragraph(instructions)
             instructions_para.paragraph_format.keep_with_next = True
         spaces = section["gap"]
         section_answers = []

         for problem in section["problems"]:
            container_table, container_cell = _start_question_block(doc)
            original_question = problem.get("question", "")
            clean_question = " ".join(original_question.split()).strip()
            suppress_question_text = problem.get("suppress_question_text")

            problem_is_multi = is_multipart_question(problem)
            graph_data = problem.get("graph") or problem.get("diagram_data")
            button_options = problem.get("button_options")
            side_by_side_layout = bool(button_options) and problem.get("side_by_side") and graph_data is not None
            has_graph = graph_data is not None
            has_spacing = spaces > 0
            should_keep_with_next = has_spacing or (tables and problem_is_multi and not side_by_side_layout) or has_graph or side_by_side_layout

            question_para = container_cell.paragraphs[0]
            question_label = f"{problem_number}."
            if suppress_question_text or not clean_question:
                question_para.text = question_label
            else:
                question_para.text = f"{question_label} {clean_question}"
            question_para.paragraph_format.keep_together = True
            question_para.paragraph_format.keep_with_next = should_keep_with_next
            if suppress_question_text:
                question_para.paragraph_format.space_after = Pt(0)
            question_para.paragraph_format.keep_with_next = should_keep_with_next

            if side_by_side_layout:
                _create_side_by_side_graph_layout(container_cell, graph_data, button_options, problem.get("units", []), problem.get("graph_doc_width"))
            elif problem_is_multi and tables:
                create_answer_table(container_cell, problem["units"], problem_number)

            if problem_is_multi:
                answers = problem["answers"]
                units = problem["units"]
                answer_parts = ["Multiple Answers:"]
                for answer, unit in zip(answers, units):
                    answer_parts.append(f"{unit}: {answer}")
                answer_str = " \n ".join(answer_parts)
                section_answers.append(f"{problem_number}. {answer_str}")
            else:
                answer = problem["answers"][0]
                unit = problem["units"][0]
                section_answers.append(f"{problem_number}. {unit}: {answer}")

            if not side_by_side_layout and graph_data:
                from utils.graph_utils import embed_graph_in_doc
                embed_graph_in_doc(container_cell, graph_data)
            effective_spaces = 0 if suppress_question_text else spaces
            for idx in range(effective_spaces):
                blank_para = container_cell.add_paragraph()
                blank_para.paragraph_format.keep_together = True
                blank_para.paragraph_format.keep_with_next = idx < effective_spaces - 1

            problem_number += 1
          
         version_answers[f"Section {section_num}: {heading}"] = section_answers
         section_num += 1
      
      answer_key.append({f"Version {doc_num}": version_answers})
     # if doc_num < number_of_docs:
         # doc.add_page_break()
          #doc.add_page_break() # ensures new page for next test regardless of current position

  if answer_key:
      doc.add_section(WD_SECTION_START.ODD_PAGE)
      doc.add_heading('Answer Key - All Versions', 1)
      for version_index, version_dict in enumerate(answer_key):
        for version_name, sections in version_dict.items():
            doc.add_heading(version_name, level=2)
            for section_name, answers in sections.items():
                doc.add_heading(section_name, level=3)
                for answer in answers:
                    doc.add_paragraph(answer)
        if version_index < len(answer_key) - 1:
            doc.add_page_break()

  file_name = f'{title}_x{number_of_docs}.docx'
  doc.save(file_name)
  try:
    os.startfile(file_name)
  except:
    pass
  print("Saving document")