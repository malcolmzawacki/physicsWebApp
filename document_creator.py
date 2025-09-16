from docx import Document
from docx.shared import Inches
import os

# Create main test document
 # Store answers as you generate

def create_documents(title: str, questions: list[dict], number_of_docs: int):
  doc = Document()
  answer_key = [] 
  for doc_num in range(1, number_of_docs + 1):
      
      # Add test header
      doc.add_heading(f'{title} {doc_num}', 0)
      doc.add_paragraph(f'Name: _____________________________________________ Date: __________________')
      
      # Generate problems here
      problem_number = 1
      version_answers = {}
      section_num = 1

      for section in questions:
         
         heading = section["heading"]
         doc.add_heading(f"Section {section_num}: {heading}", level=2)
         spaces = section["gap"]
         section_answers = []
         for problem in section["problems"]:
            clean_question = problem["question"].split(" ")
            clean_tuple = ()
            for word in clean_question:
                clean_tuple += tuple((word, " "))
            clean_question = "".join(tuple(clean_tuple))
            doc.add_paragraph(f'{problem_number}. {clean_question}')

            if "graph" in problem:
                from utils.graph_utils import embed_graph_in_doc
                embed_graph_in_doc(doc, problem["graph"])
            for i in range(spaces):
                doc.add_paragraph()
            answers = problem["answers"]
            units = problem["units"]
            answer_tuple = ()
            for answer, unit in zip(answers, units):
                unit_split = unit.split('(')
                unit = unit_split[1].replace(')','')
                answer = tuple(f"{answer} {unit},   ")
                answer_tuple += answer
            answer_str = "".join(answer_tuple)
            section_answers.append(f"{problem_number}. {answer_str}")
            problem_number += 1
          
         version_answers[f"Section {section_num}: {heading}"] = section_answers
         section_num += 1
      

      
      # Store answers for this version
      answer_key.append({f"Version {doc_num}": version_answers})
      # ensure full page gap between pages
      doc.add_page_break()


  doc.add_heading('Answer Key - All Versions', 0)
  for version_dict in answer_key:
          for version_name, sections in version_dict.items():
              doc.add_heading(version_name, level=2)
              for section_name, answers in sections.items():
                  doc.add_heading(section_name, level=3)
                  for answer in answers:
                      doc.add_paragraph(answer)
              doc.add_paragraph('')  # Space between versions

  # Save main document
  file_name = f'{title}_x{number_of_docs}.docx'
  doc.save(file_name)
  try:
    os.startfile(file_name)
  except:
    pass
  print("Saving document")

def create_and_embed_graph(doc, graph_data, filename = "temp_graph.png"):
    import matplotlib.pyplot as plt
    plt.figure(figsize=(4,4))
    plt.plot(graph_data['x'], graph_data['y'])
    plt.xlabel(graph_data['xlabel'])
    plt.ylabel(graph_data['ylabel'])
    plt.title(graph_data['title'])
    plt.grid(True)

    plt.savefig(filename, dpi=300,bbox_inches='tight')
    plt.close()

    doc.add_picture(filename, width = Inches(5))

    os.remove(filename)
    

def create_doc_from_gen(title: str, question_generator, number_of_docs: int):
  doc = Document()
  answer_key = [] 
  for doc_num in range(1, number_of_docs + 1):
      questions = question_generator()
      # Add test header
      doc.add_heading(f'{title} {doc_num}', 0)
      doc.add_paragraph(f'Name: _____________________________________________ Date: __________________')
      
      # Generate problems here
      problem_number = 1
      version_answers = {}
      section_num = 1

      for section in questions:
         
         heading = section["heading"]
         doc.add_heading(f"Section {section_num}: {heading}", level=2)
         spaces = section["gap"]
         section_answers = []
         for problem in section["problems"]:
            clean_question = problem["question"].split(" ")
            clean_tuple = ()
            for word in clean_question:
                clean_tuple += tuple((word, " "))
            clean_question = "".join(tuple(clean_tuple))
            doc.add_paragraph(f'{problem_number}. {clean_question}')

            if "graph" in problem:
                from utils.graph_utils import embed_graph_in_doc
                embed_graph_in_doc(doc, problem["graph"])
            for i in range(spaces):
                doc.add_paragraph()
            answers = problem["answers"]
            units = problem["units"]
            answer_tuple = ()
            for answer, unit in zip(answers, units):
                unit_split = unit.split('(')
                unit = unit_split[1].replace(')','')
                answer = tuple(f"{answer} {unit},   ")
                answer_tuple += answer
            answer_str = "".join(answer_tuple)
            section_answers.append(f"{problem_number}. {answer_str}")
            problem_number += 1
          
         version_answers[f"Section {section_num}: {heading}"] = section_answers
         section_num += 1
      

      
      # Store answers for this version
      answer_key.append({f"Version {doc_num}": version_answers})
      # ensure full page gap between pages
      doc.add_page_break()
      doc.add_page_break()


  doc.add_heading('Answer Key - All Versions', 0)
  for version_dict in answer_key:
          for version_name, sections in version_dict.items():
              doc.add_heading(version_name, level=2)
              for section_name, answers in sections.items():
                  doc.add_heading(section_name, level=3)
                  for answer in answers:
                      doc.add_paragraph(answer)
              doc.add_paragraph('')  # Space between versions

  # Save main document
  file_name = f'{title}_x{number_of_docs}.docx'
  doc.save(file_name)
  try:
    os.startfile(file_name)
  except:
    pass
  print("Saving document")